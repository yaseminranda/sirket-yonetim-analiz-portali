"""Backend service layer for the Contracts page: read-only viewing and receipt/invoice generation for vehicle and housing rental contracts."""
import io
from datetime import datetime
from typing import Optional

import pandas as pd
from docx import Document
from docx.shared import Pt
from sqlalchemy import text

from database import run_query


def get_all_contracts(
    search: Optional[str] = None, category: Optional[str] = None, status: Optional[str] = None
) -> pd.DataFrame:
    """Return vehicle and housing contracts unioned into one DataFrame, with optional search/category/status filters."""
    query = text(
        """
        SELECT
            ak.sozlesme_no, 'ARAC' AS kategori, ak.musteri_id, m.isim AS musteri_adi,
            ak.baslangic_tarihi, ak.bitis_tarihi, ak.total_kira AS toplam_tutar,
            ak.odenen_toplam_tutar, ak.kalan_borc, ak.sozlesme_durumu,
            a.plaka AS varlik_ozeti
        FROM araba_kiralama_sozlesmeleri ak
        LEFT JOIN musteriler m ON ak.musteri_id = m.musteri_id
        LEFT JOIN arabalar a ON ak.arac_id = a.arac_id
        UNION ALL
        SELECT
            ek.sozlesme_no, 'EV' AS kategori, ek.musteri_id, m.isim AS musteri_adi,
            ek.baslangic_tarihi, ek.bitis_tarihi, ek.total_kira AS toplam_tutar,
            ek.odenen_toplam_tutar, ek.kalan_borc, ek.sozlesme_durumu,
            (COALESCE(ap.apartman_adi, '') || ' No:' || COALESCE(d.daire_no, '')) AS varlik_ozeti
        FROM ev_kiralama_sozlesmeleri ek
        LEFT JOIN musteriler m ON ek.musteri_id = m.musteri_id
        LEFT JOIN daireler d ON ek.daire_id = d.daire_id
        LEFT JOIN apartmanlar ap ON d.apartman_id = ap.apartman_id
        ORDER BY baslangic_tarihi DESC;
        """
    )
    df = run_query(query)
    if df is None or df.empty:
        return pd.DataFrame()

    if category and category != "TÜMÜ":
        df = df[df["kategori"] == category]
    if status:
        df = df[df["sozlesme_durumu"].astype(str).str.upper().str.contains(status.upper(), na=False)]
    if search:
        s = search.lower()
        df = df[
            df["sozlesme_no"].astype(str).str.lower().str.contains(s, na=False)
            | df["musteri_adi"].astype(str).str.lower().str.contains(s, na=False)
        ]
    return df


def _table_for_category(category: str) -> str:
    """Return the DB table name for the given contract category."""
    return "ev_kiralama_sozlesmeleri" if category == "EV" else "araba_kiralama_sozlesmeleri"


def get_contract_detail(sozlesme_no: str, category: str) -> Optional[dict]:
    """Return full details of one contract (customer, vehicle/apartment, payments, drivers) as a dict."""
    table = _table_for_category(category)

    if category == "EV":
        query = text(
            f"""
            SELECT ek.*, m.isim AS musteri_adi, m.telefon AS musteri_telefon, m.email AS musteri_email,
                   m.tc_kimlik_no AS musteri_tc, ap.apartman_adi, ap.il, ap.ilce, d.daire_no, d.oda_sayisi,
                   c.ad_soyad AS calisan_adi
            FROM {table} ek
            LEFT JOIN musteriler m ON ek.musteri_id = m.musteri_id
            LEFT JOIN daireler d ON ek.daire_id = d.daire_id
            LEFT JOIN apartmanlar ap ON d.apartman_id = ap.apartman_id
            LEFT JOIN calisan c ON ek.islemi_yapan_calisan_id = c.calisan_id
            WHERE ek.sozlesme_no = :no;
            """
        )
    else:
        query = text(
            f"""
            SELECT ak.*, m.isim AS musteri_adi, m.telefon AS musteri_telefon, m.email AS musteri_email,
                   m.tc_kimlik_no AS musteri_tc, a.plaka, amar.marka_adi, am.model_adi,
                   c.ad_soyad AS calisan_adi
            FROM {table} ak
            LEFT JOIN musteriler m ON ak.musteri_id = m.musteri_id
            LEFT JOIN arabalar a ON ak.arac_id = a.arac_id
            LEFT JOIN araba_modelleri am ON a.model_id = am.model_id
            LEFT JOIN araba_markalari amar ON am.marka_id = amar.marka_id
            LEFT JOIN calisan c ON ak.islemi_yapan_calisan_id = c.calisan_id
            WHERE ak.sozlesme_no = :no;
            """
        )
    df = run_query(query, params={"no": sozlesme_no})
    if df is None or df.empty:
        return None
    detail = df.iloc[0].to_dict()
    detail["kategori"] = category

    payments_df = run_query(
        text(
            """
            SELECT odeme_id, odenen_tutar, doviz_cinsi, odenen_tutar_doviz, kur, odeme_yontemi,
                   odeme_tarihi, odeme_tipi, aciklama
            FROM odemeler WHERE sozlesme_no = :no ORDER BY odeme_tarihi DESC, odeme_id DESC;
            """
        ),
        params={"no": sozlesme_no},
    )
    detail["odemeler"] = payments_df.to_dict(orient="records") if payments_df is not None else []

    if category == "ARAC":
        drivers_df = run_query(
            text(
                """
                SELECT s.sofor_id, s.ad_soyad, s.telefon, s.email, s.tc_kimlik_no, asf.sira
                FROM arac_sozlesme_soforler asf
                JOIN soforler s ON asf.sofor_id = s.sofor_id
                WHERE asf.sozlesme_no = :no
                ORDER BY asf.sira ASC;
                """
            ),
            params={"no": sozlesme_no},
        )
        detail["soforler"] = drivers_df.to_dict(orient="records") if drivers_df is not None else []

    return detail


def generate_bulk_invoice_excel(sozlesme_no: str, category: str) -> Optional[io.BytesIO]:
    """Generate an Excel workbook listing all payments for a contract, with a totals row."""
    detail = get_contract_detail(sozlesme_no, category)
    if detail is None:
        return None
    payments = detail.get("odemeler", [])

    rows = [
        {
            "Ödeme ID": p.get("odeme_id"),
            "Sözleşme No": sozlesme_no,
            "Tarih": p.get("odeme_tarihi"),
            "Tutar (TL Karşılığı)": p.get("odenen_tutar"),
            "Döviz Cinsi": p.get("doviz_cinsi"),
            "Döviz Tutarı": p.get("odenen_tutar_doviz"),
            "Kur": p.get("kur"),
            "Ödeme Yöntemi": p.get("odeme_yontemi"),
            "İşlem Tipi": p.get("odeme_tipi"),
            "Açıklama": p.get("aciklama"),
        }
        for p in payments
    ]
    df = pd.DataFrame(
        rows,
        columns=[
            "Ödeme ID", "Sözleşme No", "Tarih", "Tutar (TL Karşılığı)", "Döviz Cinsi",
            "Döviz Tutarı", "Kur", "Ödeme Yöntemi", "İşlem Tipi", "Açıklama",
        ],
    )
    toplam = df["Tutar (TL Karşılığı)"].sum() if not df.empty else 0
    toplam_row = pd.DataFrame(
        [{"Ödeme ID": "", "Sözleşme No": "", "Tarih": "", "Tutar (TL Karşılığı)": toplam, "Döviz Cinsi": "",
          "Döviz Tutarı": "", "Kur": "", "Ödeme Yöntemi": "", "İşlem Tipi": "TOPLAM", "Açıklama": ""}]
    )
    df = pd.concat([df, toplam_row], ignore_index=True)

    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Toplu_Makbuz")
    buffer.seek(0)
    return buffer


def generate_single_payment_invoice_excel(odeme_id: int) -> Optional[io.BytesIO]:
    """Generate an Excel receipt for a single payment, re-generatable at any time."""
    df = run_query(
        text(
            """
            SELECT o.odeme_id, o.sozlesme_no, o.odenen_tutar, o.doviz_cinsi, o.odenen_tutar_doviz,
                   o.kur, o.odeme_yontemi, o.odeme_tarihi, o.odeme_tipi, o.aciklama, m.isim AS musteri_adi
            FROM odemeler o
            LEFT JOIN musteriler m ON o.musteri_id = m.musteri_id
            WHERE o.odeme_id = :id;
            """
        ),
        params={"id": odeme_id},
    )
    if df is None or df.empty:
        return None
    row = df.iloc[0].to_dict()
    data = [
        {
            "Makbuz No": f"MKB-{odeme_id}",
            "Ödeme ID": row["odeme_id"],
            "Sözleşme No": row["sozlesme_no"],
            "Müşteri": row["musteri_adi"],
            "Tarih": row["odeme_tarihi"],
            "Tutar (TL Karşılığı)": row["odenen_tutar"],
            "Döviz Cinsi": row["doviz_cinsi"],
            "Döviz Tutarı": row["odenen_tutar_doviz"],
            "Kur": row["kur"],
            "Ödeme Yöntemi": row["odeme_yontemi"],
            "İşlem Tipi": row["odeme_tipi"],
            "Açıklama": row["aciklama"],
        }
    ]
    df_out = pd.DataFrame(data)
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df_out.to_excel(writer, index=False, sheet_name="Makbuz")
    buffer.seek(0)
    return buffer


def _fmt(value) -> str:
    """Format a value for display, returning '-' for None/NaN."""
    return "-" if value is None or (isinstance(value, float) and pd.isna(value)) else str(value)


def generate_contract_docx(sozlesme_no: str, category: str) -> Optional[io.BytesIO]:
    """Generate a read-only Word document summarizing a contract's terms, parties, and payment history."""
    detail = get_contract_detail(sozlesme_no, category)
    if detail is None:
        return None

    doc = Document()

    title = doc.add_heading("KİRALAMA SÖZLEŞMESİ", level=0)
    title.alignment = 1

    doc.add_paragraph(f"Sözleşme No: {sozlesme_no}    |    Kategori: {'Araç Kiralama' if category == 'ARAC' else 'Konut (Ev) Kiralama'}")
    doc.add_paragraph(f"Belge Oluşturma Tarihi: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("1. Taraflar", level=1)
    p = doc.add_paragraph()
    p.add_run("Kiracı / Müşteri Bilgileri:\n").bold = True
    p.add_run(f"Ad Soyad: {_fmt(detail.get('musteri_adi'))}\n")
    p.add_run(f"Telefon: {_fmt(detail.get('musteri_telefon'))}\n")
    p.add_run(f"E-posta: {_fmt(detail.get('musteri_email'))}\n")
    p.add_run(f"TC Kimlik No: {_fmt(detail.get('musteri_tc'))}\n")
    p.add_run(f"\nSözleşmeyi Düzenleyen Çalışan: {_fmt(detail.get('calisan_adi'))}")

    doc.add_heading("2. Kiralanan Varlık", level=1)
    if category == "ARAC":
        doc.add_paragraph(
            f"Araç: {_fmt(detail.get('marka_adi'))} {_fmt(detail.get('model_adi'))}\n"
            f"Plaka: {_fmt(detail.get('plaka'))}"
        )
    else:
        doc.add_paragraph(
            f"Apartman/Bina: {_fmt(detail.get('apartman_adi'))}\n"
            f"Daire No: {_fmt(detail.get('daire_no'))}    Oda Sayısı: {_fmt(detail.get('oda_sayisi'))}\n"
            f"İl/İlçe: {_fmt(detail.get('il'))} / {_fmt(detail.get('ilce'))}"
        )

    doc.add_heading("3. Sözleşme Süresi ve Bedeli", level=1)
    doc.add_paragraph(f"Başlangıç Tarihi: {_fmt(detail.get('baslangic_tarihi'))}")
    doc.add_paragraph(f"Bitiş Tarihi: {_fmt(detail.get('bitis_tarihi'))}")
    doc.add_paragraph(f"Toplam Sözleşme Bedeli: ₺{float(detail.get('total_kira') or 0):,.2f}")
    doc.add_paragraph(f"Ödenen Tutar: ₺{float(detail.get('odenen_toplam_tutar') or 0):,.2f}")
    doc.add_paragraph(f"Kalan Borç: ₺{float(detail.get('kalan_borc') or 0):,.2f}")
    doc.add_paragraph(f"Sözleşme Durumu: {_fmt(detail.get('sozlesme_durumu'))}")
    if category == "EV":
        doc.add_paragraph(f"Aylık Kira: ₺{float(detail.get('aylik_kira_yrd') or 0):,.2f}")
        doc.add_paragraph(f"Depozito: ₺{float(detail.get('depozito') or 0):,.2f}")

    if category == "ARAC" and detail.get("soforler"):
        doc.add_heading("4. Ek Şoförler", level=1)
        for d in detail["soforler"]:
            doc.add_paragraph(
                f"{d.get('sira')}. Şoför: {_fmt(d.get('ad_soyad'))} | Telefon: {_fmt(d.get('telefon'))} | "
                f"TC: {_fmt(d.get('tc_kimlik_no'))}",
                style="List Bullet",
            )

    doc.add_heading("5. Ödeme Geçmişi", level=1)
    odemeler = detail.get("odemeler", [])
    if not odemeler:
        doc.add_paragraph("Bu sözleşmeye ait henüz kayıtlı bir ödeme bulunmuyor.")
    else:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Tarih", "Tutar (₺)", "Döviz", "Yöntem", "İşlem Tipi"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.font.bold = True
        for p_row in odemeler:
            cells = table.add_row().cells
            cells[0].text = _fmt(p_row.get("odeme_tarihi"))
            cells[1].text = f"{float(p_row.get('odenen_tutar') or 0):,.2f}"
            cells[2].text = _fmt(p_row.get("doviz_cinsi"))
            cells[3].text = _fmt(p_row.get("odeme_yontemi"))
            cells[4].text = _fmt(p_row.get("odeme_tipi"))

    footer = doc.add_paragraph()
    footer.add_run(
        "\nBu belge, sistemde kayıtlı sözleşme bilgilerinden otomatik olarak oluşturulmuştur ve "
        "bilgilendirme amaçlıdır."
    ).italic = True
    for run in footer.runs:
        run.font.size = Pt(9)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
